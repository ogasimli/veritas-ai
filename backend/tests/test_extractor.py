import io
import pytest
from docx import Document
from app.services.extractor import ExtractorService

def create_sample_docx() -> bytes:
    doc = Document()
    # Add a heading that should be detected as # (Heading 1)
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This is a test paragraph.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Value 1'
    table.cell(1, 1).text = 'Value 2'
    
    doc.add_paragraph('Follow up text.')
    
    # Add a level 2 heading
    doc.add_heading('Subsection 1.1', level=2)
    doc.add_paragraph('Sub-content.')
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def test_extract_markdown():
    service = ExtractorService()
    docx_bytes = create_sample_docx()
    
    markdown = service.extract_markdown(docx_bytes)
    
    # Assert order and content
    expected_parts = [
        "# Section 1",
        "This is a test paragraph.",
        "| Header 1 | Header 2 |",
        "| --- | --- |",
        "| Value 1 | Value 2 |",
        "Follow up text.",
        "## Subsection 1.1",
        "Sub-content."
    ]
    
    for part in expected_parts:
        assert part in markdown
    
    # Verify relative order (crude check)
    assert markdown.find("# Section 1") < markdown.find("Value 1")
    assert markdown.find("Value 1") < markdown.find("## Subsection 1.1")
